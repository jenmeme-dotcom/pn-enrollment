import json
import os
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT_DIR / "dist" / "documents"
OUT_FILE = OUT_DIR / "introduction-to-nursing-practical-nursing-syllabus.docx"


def load_course():
    node_bin = os.environ.get("NODE_BIN", "node")
    script = "console.log(JSON.stringify(require('./src/introNursingBuildout').introNursingCourse))"
    result = subprocess.run(
        [node_bin, "-e", script],
        cwd=ROOT_DIR,
        check=True,
        text=True,
        capture_output=True,
    )
    return json.loads(result.stdout)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E1E8", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc, course):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(course["title"])
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("12-week practical nursing course syllabus and instructor buildout")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor.from_string("526571")


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table)
    set_table_borders(table)
    set_cell_margins(table)

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_width(cell, widths[index])
        set_cell_shading(cell, "F2F4F7")
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("173F52")

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_width(cell, widths[index])
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(str(value))
    return table


def add_course_at_a_glance(doc, course):
    rows = [
        ("Course Number", course["courseNumber"]),
        ("Program", "Practical Nursing"),
        ("Length", "12 weeks"),
        ("Contact Hours", f"{course['hours']} hours"),
        ("Credits", f"{course['credits']} credits"),
        ("Assessments", "Biweekly quizzes in Weeks 2, 4, 8, and 10; midterm in Week 6; cumulative final in Week 12"),
    ]
    add_table(doc, ["Item", "Details"], rows, [2100, 7260])


def add_grading(doc, course):
    rows = [(item["title"], item["pointsPossible"]) for item in course["gradeItems"]]
    add_table(doc, ["Grade Item", "Points"], rows, [7400, 1960])
    total = sum(item["pointsPossible"] for item in course["gradeItems"])
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.add_run(f"Total possible points: {total}").bold = True


def add_assessment_calendar(doc):
    rows = [
        ("Week 2", "Quiz 1", "Weeks 1-2: nursing identity, purpose, history, and leaders"),
        ("Week 4", "Quiz 2", "Weeks 3-4: caring, communication, team roles, scope, and delegation"),
        ("Week 6", "Midterm Exam", "Weeks 1-6: nursing history, leaders, ethics, legal foundations, professionalism"),
        ("Week 8", "Quiz 3", "Weeks 7-8: culture, health equity, safety, quality, and infection prevention"),
        ("Week 10", "Quiz 4", "Weeks 9-10: nursing process, clinical judgment, teaching, and community impact"),
        ("Week 12", "Cumulative Final Exam", "Full course review and application"),
    ]
    add_table(doc, ["Week", "Assessment", "Coverage"], rows, [1400, 2400, 5560])


def add_weekly_plan(doc, course):
    for week in course["weeks"]:
        doc.add_heading(f"Week {week['week']}: {week['title']}", level=2)
        leader = doc.add_paragraph()
        leader.add_run("Nursing leaders: ").bold = True
        leader.add_run(week["nursingLeaders"])

        focus = doc.add_paragraph()
        focus.add_run("Focus: ").bold = True
        focus.add_run(week["focus"])

        doc.add_paragraph("Learning objectives:", style=None).runs[0].bold = True
        add_bullets(doc, week["objectives"])

        doc.add_paragraph("Key topics:", style=None).runs[0].bold = True
        add_bullets(doc, week["topics"])

        activity = doc.add_paragraph()
        activity.add_run("Learning activity: ").bold = True
        activity.add_run(week["activities"])

        assessment = doc.add_paragraph()
        assessment.add_run("Assessment: ").bold = True
        assessment.add_run(week["assessment"])


def build_docx():
    course = load_course()
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style_document(doc)
    add_title(doc, course)

    doc.add_heading("Course Description", level=1)
    doc.add_paragraph(course["description"])

    doc.add_heading("Course at a Glance", level=1)
    add_course_at_a_glance(doc, course)

    doc.add_heading("Purpose of Nursing", level=1)
    doc.add_paragraph(
        "Nursing exists to protect and promote health, prevent illness, relieve suffering, support healing, "
        "teach patients and families, advocate for safe care, and preserve human dignity. Earlier nursing was "
        "often informal, religious, military, or domestic caregiving. Today, nursing is a licensed, educated, "
        "accountable health profession that uses evidence, teamwork, technology, ethics, and clinical judgment."
    )
    doc.add_paragraph(
        "Practical nursing students make an impact by learning safe habits early: observing carefully, reporting "
        "changes promptly, communicating respectfully, protecting privacy, preventing infection, helping patients "
        "with basic needs, and treating every person as worthy of dignity."
    )

    doc.add_heading("Course Objectives", level=1)
    add_bullets(doc, course["objectives"])

    doc.add_heading("Required and Assigned Resources", level=1)
    add_bullets(doc, course["requiredTitles"])

    doc.add_heading("Major Ethical and Legal Topics", level=1)
    add_bullets(
        doc,
        [
            "Ethical principles: autonomy, beneficence, nonmaleficence, justice, fidelity, veracity, accountability, confidentiality, and respect for dignity.",
            "Professional boundaries, social media caution, privacy, patient rights, informed consent, refusal of care, and culturally respectful care.",
            "Legal foundations: scope of practice, student nurse limitations, delegation, supervision, negligence, malpractice, abandonment, incident reporting, mandatory reporting, and accurate documentation.",
            course["policies"]["ethicsLegal"],
        ],
    )

    doc.add_heading("Grading and Assessment", level=1)
    add_grading(doc, course)

    doc.add_heading("Quiz, Midterm, and Final Schedule", level=1)
    add_assessment_calendar(doc)

    doc.add_heading("Weekly Course Plan", level=1)
    add_weekly_plan(doc, course)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Instructor Implementation Notes", level=1)
    add_numbered(
        doc,
        [
            "Align readings to the adopted practical nursing textbook and student handbook before the course opens.",
            "Add current state board of nursing and nurse practice act references for the school location before teaching legal content.",
            "Attach program-specific attendance, grading, remediation, academic integrity, and clinical conduct policies.",
            "Set cohort-specific due dates after the cohort calendar is confirmed.",
            "Use examples from local clinical settings while protecting patient privacy and avoiding any identifying information.",
        ],
    )

    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build_docx()
